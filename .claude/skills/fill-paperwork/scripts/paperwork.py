#!/usr/bin/env python3
"""
paperwork.py — form-filling toolkit for the fill-paperwork skill.

Handles four document types from one CLI:
  - Fillable PDFs (AcroForm)  : inspect, fill-pdf
  - Flat / scanned PDFs       : inspect (text+coords), overlay-pdf
  - DOCX templates            : fill-docx
  - Any DOCX/ODT              : to-pdf (via LibreOffice)

Dependencies: PyMuPDF (fitz), python-docx, LibreOffice (soffice on PATH).
Run scripts/bootstrap.sh first if imports fail.

All subcommands print a JSON result to stdout so the agent can parse outcomes.
"""
import argparse
import json
import os
import shutil
import subprocess
import sys


def _out(obj, code=0):
    print(json.dumps(obj, indent=2, default=str))
    sys.exit(code)


def _err(msg):
    _out({"ok": False, "error": msg}, code=1)


# --------------------------------------------------------------------------
# PDF inspection
# --------------------------------------------------------------------------
def cmd_inspect(args):
    try:
        import fitz
    except ImportError:
        _err("PyMuPDF not installed. Run scripts/bootstrap.sh")

    doc = fitz.open(args.pdf)
    fields = []
    for pno, page in enumerate(doc):
        for w in page.widgets() or []:
            fields.append({
                "page": pno,
                "name": w.field_name,
                "label": w.field_label,
                "type": w.field_type_string,
                "value": w.field_value,
                "options": getattr(w, "choice_values", None),
                "rect": [round(c, 1) for c in w.rect],
            })
    result = {
        "ok": True,
        "pdf": args.pdf,
        "pages": doc.page_count,
        "fillable": len(fields) > 0,
        "field_count": len(fields),
        "fields": fields,
    }
    # For flat PDFs (no form fields), emit text blocks with coordinates so the
    # agent can decide where to overlay text.
    if not fields and args.with_text:
        blocks = []
        for pno, page in enumerate(doc):
            for b in page.get_text("blocks"):
                x0, y0, x1, y1, text = b[0], b[1], b[2], b[3], b[4]
                t = text.strip()
                if t:
                    blocks.append({
                        "page": pno,
                        "rect": [round(x0, 1), round(y0, 1), round(x1, 1), round(y1, 1)],
                        "text": t,
                    })
        result["text_blocks"] = blocks
        result["page_sizes"] = [
            [round(p.rect.width, 1), round(p.rect.height, 1)] for p in doc
        ]
    doc.close()
    _out(result)


# --------------------------------------------------------------------------
# Fill an AcroForm (fillable) PDF
# --------------------------------------------------------------------------
def cmd_fill_pdf(args):
    try:
        import fitz
    except ImportError:
        _err("PyMuPDF not installed. Run scripts/bootstrap.sh")

    with open(args.data) as f:
        data = json.load(f)  # { field_name: value, ... }

    doc = fitz.open(args.pdf)
    filled, missing, unmatched = [], [], list(data.keys())
    for page in doc:
        for w in page.widgets() or []:
            name = w.field_name
            if name in data:
                val = data[name]
                ftype = w.field_type
                if ftype == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                    on = w.on_state() or "Yes"
                    truthy = str(val).strip().lower() in ("1", "true", "yes", "on", "x", "checked")
                    w.field_value = on if truthy else "Off"
                    w.update()
                    filled.append(name)
                elif ftype == fitz.PDF_WIDGET_TYPE_RADIOBUTTON:
                    # Each widget in the group is one option. Select the one
                    # whose export (on) state matches the requested value, by
                    # raw or space-decoded ("#20" -> " ") comparison.
                    target = str(val).strip().lower()
                    on = w.on_state() or ""
                    on_dec = on.replace("#20", " ").strip().lower()
                    if target in (on.strip().lower(), on_dec):
                        w.field_value = w.on_state()
                        w.update()
                        filled.append(name)
                    else:
                        continue  # leave this option off; don't count unmatched
                else:
                    w.field_value = str(val)
                    w.update()
                    filled.append(name)
                if name in unmatched:
                    unmatched.remove(name)

    field_names = {w.field_name for p in doc for w in (p.widgets() or [])}
    for k in data:
        if k not in field_names:
            missing.append(k)

    out_path = args.out or _suffix(args.pdf, ".filled.pdf")
    doc.save(out_path, garbage=4, deflate=True)
    doc.close()
    _out({
        "ok": True,
        "out": out_path,
        "filled": sorted(set(filled)),
        "data_keys_not_in_pdf": sorted(set(missing)),
    })


# --------------------------------------------------------------------------
# Overlay text/checkmarks on a flat or scanned PDF
# --------------------------------------------------------------------------
def cmd_overlay_pdf(args):
    try:
        import fitz
    except ImportError:
        _err("PyMuPDF not installed. Run scripts/bootstrap.sh")

    with open(args.data) as f:
        items = json.load(f)  # [ {page, x, y, text, size?, font?, check?}, ... ]

    doc = fitz.open(args.pdf)
    placed = 0
    for it in items:
        page = doc[it["page"]]
        x, y = float(it["x"]), float(it["y"])
        size = float(it.get("size", 11))
        color = it.get("color", [0, 0, 0])
        if it.get("check"):
            text = "X"
        else:
            text = str(it.get("text", ""))
        page.insert_text((x, y), text, fontsize=size,
                         fontname=it.get("font", "helv"),
                         color=color)
        placed += 1

    out_path = args.out or _suffix(args.pdf, ".filled.pdf")
    doc.save(out_path, garbage=4, deflate=True)
    doc.close()
    _out({"ok": True, "out": out_path, "items_placed": placed})


# --------------------------------------------------------------------------
# Fill a DOCX template ({{placeholder}} replacement, run-aware)
# --------------------------------------------------------------------------
def cmd_fill_docx(args):
    try:
        import docx
    except ImportError:
        _err("python-docx not installed. Run scripts/bootstrap.sh")

    with open(args.data) as f:
        data = json.load(f)  # { "placeholder": "value" }

    document = docx.Document(args.docx)
    replaced = set()

    def replace_in_paragraph(p):
        full = "".join(run.text for run in p.runs)
        new = full
        for k, v in data.items():
            token = "{{" + k + "}}"
            if token in new:
                new = new.replace(token, str(v))
                replaced.add(k)
        if new != full and p.runs:
            # Collapse all text into the first run; clear the rest. Loses
            # per-run formatting mid-paragraph but is reliable for templates.
            p.runs[0].text = new
            for run in p.runs[1:]:
                run.text = ""

    for p in document.paragraphs:
        replace_in_paragraph(p)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
    for section in document.sections:
        for hf in (section.header, section.footer):
            for p in hf.paragraphs:
                replace_in_paragraph(p)

    out_path = args.out or _suffix(args.docx, ".filled.docx")
    document.save(out_path)
    _out({
        "ok": True,
        "out": out_path,
        "replaced": sorted(replaced),
        "unused_keys": sorted(set(data) - replaced),
    })


# --------------------------------------------------------------------------
# Convert DOCX/ODT -> PDF via LibreOffice
# --------------------------------------------------------------------------
def cmd_to_pdf(args):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        _err("LibreOffice (soffice) not found on PATH. Deliver the filled "
             "source file (.docx) instead, or install LibreOffice.")
    src = os.path.abspath(args.src)
    out_dir = os.path.abspath(args.out_dir or os.path.dirname(src) or ".")
    # Use an isolated, writable profile + HOME so headless conversion does not
    # collide with a locked/missing user profile.
    import tempfile
    prof = tempfile.mkdtemp(prefix="lo_profile_")
    env = dict(os.environ, HOME=prof)
    proc = subprocess.run(
        [soffice, "--headless",
         f"-env:UserInstallation=file://{prof}/cfg",
         "--convert-to", "pdf", "--outdir", out_dir, src],
        capture_output=True, text=True, timeout=180, env=env,
    )
    base = os.path.splitext(os.path.basename(src))[0] + ".pdf"
    out_path = os.path.join(out_dir, base)
    if not os.path.exists(out_path):
        _err("LibreOffice conversion failed in this environment "
             f"(stderr={proc.stderr.strip()}). The filled source file is still "
             "valid — deliver it directly, or convert on a machine with a "
             "working LibreOffice/Word install.")
    _out({"ok": True, "out": out_path})


def _suffix(path, suffix):
    base, _ = os.path.splitext(path)
    return base + suffix


def main():
    ap = argparse.ArgumentParser(description="Form-filling toolkit")
    sub = ap.add_subparsers(dest="cmd", required=True)

    p = sub.add_parser("inspect", help="List form fields, or text blocks if flat")
    p.add_argument("pdf")
    p.add_argument("--with-text", action="store_true",
                   help="For flat PDFs, also dump text blocks with coordinates")
    p.set_defaults(func=cmd_inspect)

    p = sub.add_parser("fill-pdf", help="Fill an AcroForm PDF from a JSON map")
    p.add_argument("pdf")
    p.add_argument("data", help="JSON file: {field_name: value}")
    p.add_argument("--out")
    p.set_defaults(func=cmd_fill_pdf)

    p = sub.add_parser("overlay-pdf", help="Overlay text on a flat/scanned PDF")
    p.add_argument("pdf")
    p.add_argument("data", help="JSON file: [{page,x,y,text,size?,check?}]")
    p.add_argument("--out")
    p.set_defaults(func=cmd_overlay_pdf)

    p = sub.add_parser("fill-docx", help="Replace {{placeholders}} in a DOCX")
    p.add_argument("docx")
    p.add_argument("data", help="JSON file: {placeholder: value}")
    p.add_argument("--out")
    p.set_defaults(func=cmd_fill_docx)

    p = sub.add_parser("to-pdf", help="Convert DOCX/ODT to PDF (LibreOffice)")
    p.add_argument("src")
    p.add_argument("--out-dir")
    p.set_defaults(func=cmd_to_pdf)

    args = ap.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
